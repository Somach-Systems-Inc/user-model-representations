"""Build the submission write-up as .docx from WRITEUP.md (a small markdown subset:
'# ', '## ', '### ' headings; '- ' bullets; '| a | b |' tables with a '|---|' rule;
'![caption](path.png)' images; blank-line paragraphs; **bold** and *italic* inline).
Then upload via the Drive MCP as a Google Doc.

    uv run python src/build_writeup.py WRITEUP.md out/writeup.docx
"""
import re, sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

src, out = sys.argv[1], sys.argv[2]
doc = Document()
st = doc.styles["Normal"]; st.font.name = "Arial"; st.font.size = Pt(10.5)
for s in ("Heading 1", "Heading 2", "Heading 3"):
    doc.styles[s].font.name = "Arial"

def add_inline(par, text):
    for tok in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text):
        if not tok: continue
        if tok.startswith("**"): r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("*"): r = par.add_run(tok[1:-1]); r.italic = True
        elif tok.startswith("`"): r = par.add_run(tok[1:-1]); r.font.name = "Courier New"
        else: par.add_run(tok)

lines = open(src, encoding="utf-8").read().splitlines()
i = 0; para = []
def flush():
    global para
    if para:
        p = doc.add_paragraph(); add_inline(p, " ".join(para)); para = []
while i < len(lines):
    ln = lines[i]
    if ln.startswith("# "): flush(); doc.add_heading(ln[2:].replace("`",""), 0)
    elif ln.startswith("## "): flush(); doc.add_heading(ln[3:].replace("`",""), 1)
    elif ln.startswith("### "): flush(); doc.add_heading(ln[4:].replace("`",""), 2)
    elif ln.startswith("- "): flush(); p = doc.add_paragraph(style="List Bullet"); add_inline(p, ln[2:])
    elif ln.startswith("!["):
        flush(); m = re.match(r"!\[(.*?)\]\((.*?)\)", ln)
        doc.add_picture(m.group(2), width=Inches(6.2))
        cap = doc.add_paragraph(); add_inline(cap, m.group(1)); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs: r.font.size = Pt(9); r.italic = True
    elif ln.startswith("|"):
        flush(); rows = []
        while i < len(lines) and lines[i].startswith("|"):
            if not re.match(r"^\|\s*-", lines[i]): rows.append([c.strip() for c in lines[i].strip("|").split("|")])
            i += 1
        t = doc.add_table(rows=len(rows), cols=len(rows[0])); t.style = "Light Grid Accent 1"
        for r, row in enumerate(rows):
            for c, cell in enumerate(row):
                t.cell(r, c).text = ""; add_inline(t.cell(r, c).paragraphs[0], cell)
                for run in t.cell(r, c).paragraphs[0].runs: run.font.size = Pt(9); run.bold = run.bold or r == 0
        continue
    elif ln.strip() == "": flush()
    else: para.append(ln.strip())
    i += 1
flush(); doc.save(out); print("saved", out)
